from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict


def build_docx(process: Dict, author: str | None, source_name: str | None, out_path: str) -> str:
    doc = Document()

    # Title Page
    title = process.get("title", "Process Document")
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}
").font.size = Pt(11)
    if author:
        meta.add_run(f"Author: {author}
").font.size = Pt(11)
    if source_name:
        meta.add_run(f"Source: {source_name}
").font.size = Pt(11)

    doc.add_page_break()

    # Overview
    doc.add_heading("Overview", level=1)
    if process.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(process["summary"])    

    if process.get("purpose"):
        doc.add_heading("Purpose", level=2)
        doc.add_paragraph(process["purpose"])    

    if process.get("scope"):
        doc.add_heading("Scope", level=2)
        doc.add_paragraph(process["scope"])    

    roles = process.get("roles", [])
    if roles:
        doc.add_heading("Roles", level=2)
        for r in roles:
            doc.add_paragraph(f"• {r}")

    tools = process.get("tools", [])
    if tools:
        doc.add_heading("Tools", level=2)
        for t in tools:
            doc.add_paragraph(f"• {t}")

    # Steps
    steps = process.get("steps", [])
    if steps:
        doc.add_heading("End-to-End Process", level=1)
        for s in steps:
            n = s.get("number")
            action = s.get("action") or "Step"
            details = s.get("details")
            role = s.get("role")
            ts = s.get("tools") or []

            doc.add_heading(f"Step {n}: {action}", level=2)
            if details:
                doc.add_paragraph(details)
            if role:
                doc.add_paragraph(f"Role Responsible: {role}")
            if ts:
                doc.add_paragraph("Tools:")
                for t in ts:
                    doc.add_paragraph(f"• {t}")

    # Decisions
    decisions = process.get("decisions", [])
    if decisions:
        doc.add_heading("Decision Points", level=1)
        for d in decisions:
            cond = d.get("condition")
            yes = d.get("path_yes")
            no = d.get("path_no")
            doc.add_paragraph(f"If: {cond}")
            if yes:
                doc.add_paragraph(f"Yes → {yes}")
            if no:
                doc.add_paragraph(f"No → {no}")

    # Notes
    notes = process.get("notes", [])
    if notes:
        doc.add_heading("Additional Notes", level=1)
        for n in notes:
            doc.add_paragraph(f"• {n}")

    doc.save(out_path)
    return out_path
